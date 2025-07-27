
import os
import re
import time
import logging
import hashlib
import requests # Kept for Groq/other HTTP needs, not for Google Translate anymore
from docx import Document

from docx.oxml import OxmlElement #, parse_xml # parse_xml not strictly used
from docx.oxml.ns import qn
from groq import Groq
import google.generativeai as genai # For Gemini API

logger = logging.getLogger(__name__)

# os.environ['GEMINI_API_KEY']=""
# os.environ["GROQ_API_KEY"] = ""


logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(name)s: %(message)s")
translation_cache = {}

class TranslationMetrics:
    def __init__(self):
        self._api_calls_total_session = 0
        self._api_calls_current_minute = 0
        self.cache_hits = 0
        self.retries = 0
        self._last_reset_minute_tracking = time.time()
        self._rpm_limit = 28 # General limit, specific APIs might have different (e.g. QPM for Gemini)
        
    @property
    def api_calls_this_minute(self):
        current_time = time.time()
        if current_time - self._last_reset_minute_tracking >= 60:
            self._api_calls_current_minute = 0
            self._last_reset_minute_tracking = current_time
        return self._api_calls_current_minute
    
    def log_api_call(self):
        self._api_calls_current_minute += 1
        self._api_calls_total_session += 1
        # Adjust warning threshold based on typical API limits (e.g., Gemini free tier is ~60 QPM)
        # For now, keeping the general warning.
        if self._api_calls_current_minute >= self._rpm_limit - 5: 
            logger.warning(f"High API call count in current minute: {self._api_calls_current_minute}")

    def get_total_api_calls_session(self):
        return self._api_calls_total_session

def wait_if_needed(metrics: TranslationMetrics, specific_rpm_limit=None):
    """Wait if we're approaching rate limits for the current minute."""
    limit_to_check = specific_rpm_limit if specific_rpm_limit is not None else metrics._rpm_limit
    if metrics.api_calls_this_minute >= limit_to_check:
        wait_time = 60 - (time.time() - metrics._last_reset_minute_tracking)
        if wait_time > 0:
            logger.info(f"Rate limit approaching ({metrics.api_calls_this_minute} calls vs limit {limit_to_check}), waiting {wait_time:.2f} seconds")
            time.sleep(wait_time)
            # After waiting, the counter for the minute should reset naturally or on next check
            return True
    return False

def is_bullet_point(text):
    bullet_patterns = [
        r'^\s*[\u2022\u2023\u2043\u204C\u204D\u2219\u25AA\u25CF\u25E6\u29BE\u29BF]\s+',
        r'^\s*[-–—•]\s+', r'^\s*\d+\.\s+', r'^\s*[a-zA-Z]\.\s+',
        r'^\s*\(\d+\)\s+', r'^\s*\[\d+\]\s+', r'^\s*\（\d+\）\s+'
    ]
    return any(re.match(pattern, text.strip()) for pattern in bullet_patterns)

def extract_bullet_marker(text):
    patterns = [
        r'^(\s*[\u2022\u2023\u2043\u204C\u204D\u2219\u25AA\u25CF\u25E6\u29BE\u29BF]\s+)',
        r'^(\s*[-–—•]\s+)', r'^(\s*\d+\.\s+)', r'^(\s*[a-zA-Z]\.\s+)',
        r'^(\s*\(\d+\)\s+)', r'^(\s*\[\d+\]\s+)', r'^(\s*\（\d+\）\s+)'
    ]
    text_stripped = text.strip()
    for pattern in patterns:
        match = re.match(pattern, text_stripped)
        if match:
            original_match_region = text[:len(match.group(1))]
            return original_match_region
    return ""

def _split_text_for_translation(text, max_length=15000): # Gemini can handle larger chunks
    """
    Splits text into chunks smaller than max_length (in characters),
    trying to preserve sentences or at least split at spaces.
    Gemini models usually have large context windows (e.g., 32k tokens for gemini-pro, 
    1M for 1.5 pro, flash is 1M). Max length here is for a single text item passed in a batch.
    This might be more relevant if a single paragraph itself is extremely long.
    """
    if len(text) <= max_length:
        return [text]
    logger.warning(f"Text item of length {len(text)} exceeds {max_length} and will be split.")
    chunks = []
    current_pos = 0
    while current_pos < len(text):
        end_pos = min(current_pos + max_length, len(text))
        if end_pos == len(text):
            chunks.append(text[current_pos:])
            break
        
        split_at = -1
        for delimiter in ['. ', '! ', '? ', '\n', '.']: # Added standalone period
            temp_split_at = text.rfind(delimiter, current_pos, end_pos)
            if temp_split_at != -1: # Potential split point
                # Ensure it's not part of a number e.g., "item 1.2"
                if delimiter == '.' and temp_split_at > current_pos and text[temp_split_at-1].isdigit() and \
                   temp_split_at + 1 < len(text) and text[temp_split_at+1].isdigit():
                    continue # Skip if it's like "1.2"
                split_at = max(split_at, temp_split_at + len(delimiter) -1)
        
        if split_at == -1: 
            split_at = text.rfind(' ', current_pos, end_pos)
        
        if split_at == -1 or split_at <= current_pos : 
            split_at = end_pos -1 
        
        chunks.append(text[current_pos : split_at + 1].strip())
        current_pos = split_at + 1
        while current_pos < len(text) and text[current_pos].isspace():
            current_pos +=1
            
    return [c for c in chunks if c.strip()]


def translation_using_gemini(texts, target_language='hi', tone='professional', max_retries=3, model_name="gemini-2.0-flash"):
    """
    Translate texts using Google Gemini API.
    Handles batching and retries.
    """
    local_metrics = TranslationMetrics()

    if not texts: return []
    input_texts = [texts] if not isinstance(texts, list) else texts
    if all(not text or not text.strip() for text in input_texts):
        return [""] * len(input_texts)

    gemini_api_key = os.environ.get("GEMINI_API_KEY")
    if not gemini_api_key:
        logger.error("GEMINI_API_KEY not found in environment variables. Cannot use Gemini engine.")
        raise RuntimeError("GEMINI_API_KEY not found in environment variables.")

    try:
        genai.configure(api_key=gemini_api_key)
        model = genai.GenerativeModel(model_name)
    except Exception as e:
        logger.error(f"Failed to configure Gemini API or model: {e}", exc_info=True)
        return input_texts

    texts_tuple = tuple(input_texts) # For caching
    cache_key = hashlib.md5(f"gemini_{target_language}_{tone}_{model_name}_{texts_tuple}".encode('utf-8')).hexdigest()
    if cache_key in translation_cache:
        local_metrics.cache_hits +=1
        logger.info(f"Cache hit for Gemini translation batch. Returning cached result.")
        return translation_cache[cache_key]

    
    processed_input_texts = []
    for text_item in input_texts:
        processed_input_texts.extend(_split_text_for_translation(text_item))
    
    # If splitting resulted in more texts than original, this logic needs adjustment
    # For now, assume _split_text_for_translation is mainly for *very* long individual items
    # and the number of items in `input_texts` is what we expect back.
    # This simplified example will translate each text item separately if batch prompting is complex.
    # A better approach is batch prompting:
    
    prompt_parts = [
        f"You are an expert linguist and translator. Translate the following text items from their original language into {target_language}.",
        f"Maintain a {tone} tone for all translations.",
        "Preserve all original formatting cues like list markers (e.g., '1.', 'a.', '*'), indentation, and special characters to the best of your ability in the translated text.",
        "Do not translate numbers, or highly technical terms and acronyms unless essential for contextual meaning in the target language.",
        "For each item provided between <ITEM> and </ITEM> tags, provide ONLY its direct translation, also enclosed in <ITEM> and </ITEM> tags.",
        "Ensure the number of output <ITEM> blocks matches the number of input <ITEM> blocks.",
        "Example Request:",
        "<ITEM>Hello world.</ITEM>",
        "<ITEM>This is a test.</ITEM>",
        "Example Response (if translating to French):",
        "<ITEM>Bonjour le monde.</ITEM>",
        "<ITEM>Ceci est un test.</ITEM>",
        "\nNow, translate the following texts:",
    ]
    
    if not processed_input_texts: # Should have been caught earlier
        return [""] * len(input_texts)

    for i, text_item in enumerate(input_texts): # Use original input_texts for item construction
        prompt_parts.append(f"<ITEM>{text_item}</ITEM>")
    
    full_prompt = "\n".join(prompt_parts)

    api_response_content = ""
    success = False

    # Gemini safety settings 
    safety_settings = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    ]
    # Generation Config
    generation_config = genai.types.GenerationConfig(
        # candidate_count=1, # Default is 1
        # stop_sequences=['</ITEM>'], # Could be useful but might also truncate valid content
        # max_output_tokens=... # Set based on expected output length
        temperature=0.25, # Lower for more deterministic translation
    )


    for attempt in range(max_retries + 1):
        logger.info(f"Attempting Gemini translation (Attempt {attempt + 1}/{max_retries + 1}) for {len(input_texts)} items.")
        try:
            response = model.generate_content(
                full_prompt,
                generation_config=generation_config,
                safety_settings=safety_settings
                )
            
            if response.prompt_feedback and response.prompt_feedback.block_reason:
                logger.error(f"Gemini prompt blocked. Reason: {response.prompt_feedback.block_reason}")
                # If specific safety ratings are available:
                for rating in response.prompt_feedback.safety_ratings:
                    logger.error(f"  Safety Rating: {rating.category} - {rating.probability}")
                if attempt < max_retries: 
                    time.sleep(5) # Wait before retrying if blocked, maybe adjust prompt or settings
                    continue 
                else: # Final attempt blocked
                    raise Exception(f"Gemini prompt blocked after retries: {response.prompt_feedback.block_reason}")


            if response.parts:
                 api_response_content = "".join(part.text for part in response.parts if hasattr(part, 'text'))
            elif hasattr(response, 'text') and response.text: # Older API or simpler response
                 api_response_content = response.text
            else: # No text part found
                logger.warning(f"Gemini response has no text part. Response: {response}")
                api_response_content = "" # Ensure it's empty

            api_response_content = api_response_content.strip()
            success = True 
            logger.info(f"Gemini translation API call successful on attempt {attempt + 1}")
            break # Exit retry loop on success
               
        except Exception as e:
            logger.error(f"Error during Gemini translation attempt {attempt + 1}: {e}", exc_info=True)
            local_metrics.retries += 1
            # Basic check for common retryable errors (e.g., rate limits, server errors)
            if "429" in str(e) or "rate limit" in str(e).lower() or "500" in str(e) or "503" in str(e):
                wait_time = min(2 ** attempt * 2, 60) 
                logger.info(f"Gemini API error (possibly retryable), waiting {wait_time} seconds before retry {attempt + 2}")
                time.sleep(wait_time)
            elif attempt < max_retries:
                time.sleep(min(2 ** attempt, 30)) 
            else: 
                logger.error(f"Gemini translation failed after {max_retries + 1} attempts.")
                break 

    processed_results = []
    if success and api_response_content:
        item_pattern = re.compile(r'<ITEM>(.*?)</ITEM>', re.DOTALL)
        found_items = item_pattern.findall(api_response_content)
        
        if len(found_items) == len(input_texts):
            processed_results = [item.strip() for item in found_items] # Strip individual items
            logger.info(f"Successfully extracted {len(found_items)} translated items from Gemini batch response.")
        else:
            logger.warning(
                f"Marker mismatch in Gemini response. Expected {len(input_texts)}, got {len(found_items)}. "
                f"Response: '{api_response_content[:300]}...'"
            )
            # Fallback strategy:
            if found_items: # If some items were found, use them
                processed_results = [item.strip() for item in found_items]
                # Pad with originals if fewer items returned
                while len(processed_results) < len(input_texts):
                    original_idx = len(processed_results)
                    logger.warning(f"Padding missing Gemini translation for item {original_idx} with original.")
                    processed_results.append(input_texts[original_idx])
                # Truncate if too many items returned (less likely but possible)
                if len(processed_results) > len(input_texts):
                    logger.warning("Gemini returned more items than expected, truncating.")
                    processed_results = processed_results[:len(input_texts)]
            else: # No items found with markers
                logger.error("No <ITEM> markers found in Gemini response. Using original texts for this batch.")
                processed_results = input_texts
                success = False 
    else:
        logger.error("Gemini translation failed or produced empty content. Using original texts for this batch.")
        processed_results = input_texts
        success = False

    if success:
        translation_cache[cache_key] = processed_results
        logger.info(f"Gemini translation batch processed successfully.")
    else:
        logger.warning(f"Gemini translation batch finished with errors/fallbacks.")
            
    return processed_results


def translation_using_groq(texts, target_language='hi', tone='professional', max_retries=2):
    """Translate texts using Groq API with smart batching for paragraphs."""
    local_metrics = TranslationMetrics()
    
    if not texts or all(not text for text in texts):
        return [""] * len(texts) if isinstance(texts, list) else [""]
    
    input_texts = [texts] if not isinstance(texts, list) else texts
        
    is_single_mode = len(input_texts) == 1 # This logic is slightly different from Gemini batching
    batch_mode_text = "SINGLE" if is_single_mode else f"BATCH({len(input_texts)})"
    # logger.info(f"Starting Groq {batch_mode_text} item translation for text: '{str(input_texts[0])[:30]}...'")

    texts_tuple = tuple(input_texts)
    cache_key = hashlib.md5(f"groq_{target_language}_{tone}_{texts_tuple}".encode('utf-8')).hexdigest()
    if cache_key in translation_cache:
        local_metrics.cache_hits += 1
        logger.info(f"Cache hit for Groq translation batch. Returning cached result.")
        return translation_cache[cache_key]
        
    try:
        api_key = os.environ.get("GROQ_API_KEY")
        if not api_key:
            logger.error("GROQ_API_KEY not found in environment variables.")
            raise RuntimeError("GROQ_API_KEY not found in environment variables.")
        client = Groq(api_key=api_key)

        system_prompt = (
            f"You are a precise and powerful translator. Translate the provided text(s) to {target_language}.\n"
            f"Maintain a {tone} tone.\n"
            f"Preserve all original formatting, numbering, spacing around special characters, and technical terms as much as possible.\n"
            f"Do not translate proper nouns, code, or highly specific technical jargon unless it has a well-known equivalent in {target_language}.\n"
        )
        
        prompt_content_parts = []
        if is_single_mode:
            system_prompt += "Translate the single text accurately. Output ONLY the translated text."
            prompt_content_parts.append(input_texts[0])
        else: # Batch mode for Groq
            system_prompt += (
                "The texts are provided in <ITEM>...</ITEM> blocks. "
                "For each input <ITEM>, provide its translation also in an <ITEM>...</ITEM> block. "
                "Ensure the output contains the same number of <ITEM> blocks as the input."
            )
            for text_item in input_texts:
                prompt_content_parts.append(f"<ITEM>{str(text_item)}</ITEM>")
        
        prompt_content = "\n".join(prompt_content_parts)
            
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt_content}
        ]

        api_response_content = ""
        success = False
        
        for attempt in range(max_retries + 1):
            logger.info(f"Attempting Groq translation (Attempt {attempt + 1}/{max_retries + 1}) for {len(input_texts)} items.")
            try:
                response = client.chat.completions.create(
                    messages=messages,
                    model="llama3-70b-8192", # Or other preferred Groq model
                    temperature=0.2, 
                )
                api_response_content = response.choices[0].message.content.strip()
                success = True 
                logger.info(f"Groq translation API call successful on attempt {attempt + 1}")
                break
                   
            except Exception as e:
                logger.error(f"Error during Groq translation attempt {attempt + 1}: {e}", exc_info=True)
                local_metrics.retries += 1
                if "429" in str(e) or "rate limit" in str(e).lower():
                    wait_time = min(2 ** attempt * 2, 60) 
                    logger.info(f"Groq rate limit hit, waiting {wait_time} seconds before retry {attempt + 2}")
                    time.sleep(wait_time)
                elif attempt < max_retries:
                    time.sleep(min(2 ** attempt, 30)) 
                else: 
                    logger.error(f"Groq translation failed after {max_retries + 1} attempts.")
                    break 

        processed_results = []
        if success and api_response_content:
            if is_single_mode:
                processed_results = [api_response_content]
            else: # Batch mode parsing
                item_pattern = re.compile(r'<ITEM>(.*?)</ITEM>', re.DOTALL)
                found_items = item_pattern.findall(api_response_content)
                
                if len(found_items) == len(input_texts):
                    processed_results = [item.strip() for item in found_items]
                    logger.info(f"Successfully extracted {len(found_items)} translated items from Groq batch response.")
                else:
                    logger.warning(
                        f"Marker mismatch in Groq batch response. Expected {len(input_texts)}, got {len(found_items)}. "
                        f"Response: '{api_response_content[:200]}...'"
                    )
                    if found_items:
                        processed_results = [item.strip() for item in found_items]
                        while len(processed_results) < len(input_texts):
                            processed_results.append(input_texts[len(processed_results)]) # Pad with original
                        if len(processed_results) > len(input_texts):
                            processed_results = processed_results[:len(input_texts)] # Truncate
                    else:
                        logger.error("No <ITEM> markers found in Groq response. Using original texts.")
                        processed_results = input_texts
                        success = False
        else:
            logger.error("Groq translation failed or produced empty content. Using original texts.")
            processed_results = input_texts
            success = False

        if success:
            translation_cache[cache_key] = processed_results
            logger.info(f"Groq {batch_mode_text} translation finished successfully.")
        else:
            logger.warning(f"Groq {batch_mode_text} translation finished with errors/fallbacks.")
            
        return processed_results

    except Exception as e:
        logger.error(f"Fatal error during Groq translation: {e}", exc_info=True)
        return texts

# --- Summarization Function (Groq) ---
def summarize_text_groq(full_text, target_language='hi', max_length=200, tone='professional'):
    if not full_text or not full_text.strip():
        logging.warning("Summarization attempt with empty text.")
        return "Could not generate summary: Input text is empty."

    logging.info(f"Requesting Groq summary in {target_language} for text ({len(full_text)} chars), max_length={max_length} words.")
    try:
        api_key = os.environ.get("GROQ_API_KEY")
        if not api_key:
            logger.error("GROQ_API_KEY not found for summarization.")
            return "Error: GROQ_API_KEY not configured."
        client = Groq(api_key=api_key)
    
        system_prompt = (
            f"You are a helpful assistant skilled in summarizing long texts. "
            f"Analyze the following English text and provide a concise summary in {target_language}. "
            f"The summary should capture the main points and be approximately {max_length} words long. "
            f"Focus on clarity and accuracy in {target_language}. The tone of the language should be {tone}. "
            f"Do not add any introductory phrases like 'Here is the summary:'."
        )
        
        max_input_chars = 20000 
        if len(full_text) > max_input_chars:
             logging.warning(f"Input text ({len(full_text)} chars) exceeds limit ({max_input_chars}), truncating for summary.")
             full_text = full_text[:max_input_chars] + "... [truncated]"

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Please summarize the following text:\n{full_text}"}
        ]
        response = client.chat.completions.create(
            messages=messages, model="llama3-70b-8192", temperature=0.3, 
            max_tokens=max_length * 5 
        ).choices[0].message.content.strip()
        logging.info(f"Groq summary generated successfully ({len(response)} chars in {target_language})." )
        return response
    except Exception as e:
        logging.error(f"Groq summarization failed: {e}", exc_info=True)
        return f"Error generating summary: An API or processing error occurred."

def get_paragraph_direction(text):
    if not text: return False
    rtl_ranges = [
        (0x0600, 0x06FF), (0x0750, 0x077F), (0x08A0, 0x08FF), (0xFB50, 0xFDFF),
        (0xFE70, 0xFEFF), (0x0590, 0x05FF), (0x0700, 0x074F), (0x0780, 0x07BF),
        (0x0840, 0x085F) 
    ]
    rtl_count = 0; ltr_count = 0
    for char in text:
        code = ord(char)
        if any(start <= code <= end for start, end in rtl_ranges): rtl_count += 1
        elif char.isalpha(): ltr_count += 1
    total_alpha_numeric = rtl_count + ltr_count
    if total_alpha_numeric == 0: return False
    return (rtl_count / total_alpha_numeric) > 0.4 

def is_rtl_language(lang_code):
    rtl_langs = {'ar', 'arc', 'dv', 'fa', 'ha', 'he', 'khw', 'ks', 'ku', 'ps', 'ur', 'yi'}
    return lang_code.lower().split('-')[0] in rtl_langs

def is_numbered_list_paragraph(para):
    try:
        return para._element.find('.//w:numPr', namespaces=para._element.nsmap) is not None
    except AttributeError: return False

def extract_list_marker(text):
    patterns = [
        r'^(\s*\.\d+\s+)', r'^(\s*\d+\.\s+)', r'^(\s*\(\d+\)\s+)', r'^(\s*\[\d+\]\s+)',
        r'^(\s*[A-Za-z]\.\s+)', r'^(\s*\([A-Za-z]\)\s+)', r'^(\s*[ivxclmdIVXCLMD]+\.\s+)',
        r'^(\s*\([ivxclmdIVXCLMD]+\)\s+)',
        r'^(\s*[\u2022\u2023\u2043\u204C\u204D\u2219\u25AA\u25CF\u25E6\u29BE\u29BF•◦▪-]\s+)'
    ]
    stripped_text = text.lstrip()
    for pattern in patterns:
        match = re.match(pattern, stripped_text)
        if match:
            original_leading_spaces = text[:len(text) - len(stripped_text)]
            return original_leading_spaces + match.group(1)
    return None

def process_translated_text(original_text, translated_text):
    original_marker = extract_list_marker(original_text)
    if original_marker:
        translated_marker = extract_list_marker(translated_text)
        translated_text_content = translated_text[len(translated_marker):] if translated_marker else translated_text.lstrip()
        return original_marker + translated_text_content
    return translated_text

def translate_docx_advanced(input_path, output_path, target_language='hi', engine='gemini', tone='professional', first_page_only=False, progress_callback=None, force_rtl=False):
    logger.info(f"Starting DOCX translation: {input_path} -> {output_path} (lang={target_language}, engine={engine}, force_rtl={force_rtl})")

    try:
        doc = Document(input_path)
    except Exception as e:
        logger.error(f"Failed to open input DOCX file: {input_path} - {e}", exc_info=True)
        return False
        
    doc_translation_metrics = TranslationMetrics() 
    target_is_rtl_lang = is_rtl_language(target_language)
    
    logger.info(f"Setting document default direction based on target language ({'RTL' if target_is_rtl_lang or force_rtl else 'LTR'})")
    for section in doc.sections:
        try:
            sectPr = section._sectPr
            if sectPr is not None: 
                bidi_tag_list = sectPr.xpath('./w:bidi')
                if target_is_rtl_lang or force_rtl:
                    if not bidi_tag_list: bidi_tag = OxmlElement('w:bidi'); sectPr.append(bidi_tag)
                    else: bidi_tag = bidi_tag_list[0]
                    bidi_tag.set(qn('w:val'), "true") 
                else:
                    if bidi_tag_list: sectPr.remove(bidi_tag_list[0])
        except Exception as e: 
            logger.warning(f"Unable to set section-level direction: {e}", exc_info=True)

    all_paragraphs_with_runs = {} 
    body_paras, header_footer_paras, table_paras, shape_paras = [], [], [], []

    logger.info("Collecting and categorizing document elements...")
    for para_container_sources in [
        doc.paragraphs,
        (p for s in doc.sections for hf_type in (s.header, s.footer) if hf_type for p in hf_type.paragraphs),
        (p for t in doc.tables for r in t.rows for c in r.cells for p in c.paragraphs),
        (p for sh in doc.inline_shapes if hasattr(sh, 'text_frame') and sh.text_frame and sh.text_frame.text.strip() for p in sh.text_frame.paragraphs)
    ]:
        for para in para_container_sources:
            if para.text.strip():
                # Categorize based on source if needed, for now add to one list
                # For simplicity, collecting all pars here. If separate batching for headers is needed, adjust.
                if para not in all_paragraphs_with_runs: # Avoid duplicates if somehow processed from different sources
                    # Determine category (simplified)
                    if any(para in lst for lst in [header_footer_paras, table_paras, shape_paras]):
                        pass # Already categorized by earlier specific loops if they were separate
                    elif para in doc.paragraphs: # Crude check, better would be to know origin
                         body_paras.append(para)
                    # This simplified collection might miscategorize if not careful with generator order.
                    # Explicit loops as before are safer for categorization. Reverting to explicit:

    # --- Reverted to Explicit Loops for Clear Categorization & Run Collection ---
    for para in doc.paragraphs:
        if para.text.strip():
            body_paras.append(para)
            all_paragraphs_with_runs[para] = [(run, run.text) for run in para.runs if run.text.strip()]
    for section in doc.sections:
        for LTYPE in ['header', 'footer']:
            container = getattr(section, LTYPE, None)
            if container:
                for para in container.paragraphs:
                    if para.text.strip():
                        header_footer_paras.append(para)
                        all_paragraphs_with_runs[para] = [(run, run.text) for run in para.runs if run.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        table_paras.append(para)
                        all_paragraphs_with_runs[para] = [(run, run.text) for run in para.runs if run.text.strip()]
    for shape in doc.inline_shapes:
        try:
            if hasattr(shape, 'text_frame') and shape.text_frame and shape.text_frame.text.strip():
                for para in shape.text_frame.paragraphs:
                    if para.text.strip():
                        shape_paras.append(para)
                        all_paragraphs_with_runs[para] = [(run, run.text) for run in para.runs if run.text.strip()]
        except AttributeError: 
            logger.debug(f"Inline shape type {type(shape)} does not have text_frame or no text.")
        except Exception as e:
            logger.warning(f"Error processing inline_shape text: {e}", exc_info=True)
    # --- End Reverted Loops ---

    all_paragraph_objects = body_paras + header_footer_paras + table_paras + shape_paras
    total_paragraphs_to_process = len(all_paragraph_objects)
    logger.info(f"Collected {total_paragraphs_to_process} non-empty paragraphs for translation.")

    if total_paragraphs_to_process == 0:
        logger.info("No text found to translate. Saving document as is.")
        doc.save(output_path); 
        if progress_callback: progress_callback(1,1)
        return True

    def create_translation_batches(paragraphs_list, max_paras_per_batch=10, max_chars_per_batch=10000): # Adjusted for LLMs
        batches = []; current_batch_paras = []; current_batch_chars = 0
        for para in paragraphs_list:
            para_text_len = len(para.text)
            if not para.text.strip(): continue
            if current_batch_paras and \
               (len(current_batch_paras) >= max_paras_per_batch or \
                current_batch_chars + para_text_len > max_chars_per_batch):
                batches.append(current_batch_paras)
                current_batch_paras = []; current_batch_chars = 0
            current_batch_paras.append(para)
            current_batch_chars += para_text_len
        if current_batch_paras: batches.append(current_batch_paras)
        logger.info(f"Created {len(batches)} batches for {len(paragraphs_list)} paragraphs.")
        return batches

    # Gemini can handle larger batches, Groq similar.
    # Max paras per batch can be higher for LLMs if they handle item markers well.
    # Max chars depends on model context window and how much overhead the prompt adds.
    paras_per_batch_limit = 15 if engine in ['gemini', 'groq'] else 5 
    chars_per_batch_limit = 20000 if engine in ['gemini', 'groq'] else 5000

    paragraph_batches_for_api = create_translation_batches(
        all_paragraph_objects, 
        max_paras_per_batch=paras_per_batch_limit,
        max_chars_per_batch=chars_per_batch_limit
    )

    logger.info(f"Starting translation for {len(paragraph_batches_for_api)} batches using engine: {engine}.")
    
    processed_paragraph_count = 0
    total_api_batches = len(paragraph_batches_for_api)

    # Determine engine-specific rate limits (RPM/QPM)
    
    engine_qpm_limit = 50 if engine == 'gemini' else (doc_translation_metrics._rpm_limit) # Conservative for Gemini

    for batch_idx, para_batch_list in enumerate(paragraph_batches_for_api):
        if wait_if_needed(doc_translation_metrics, specific_rpm_limit=engine_qpm_limit):
             logger.info(f"Resuming after rate limit wait. Batch {batch_idx+1}/{total_api_batches}")
        
        original_texts_in_batch = [p.text for p in para_batch_list]
        if not any(text.strip() for text in original_texts_in_batch):
            processed_paragraph_count += len(para_batch_list)
            if progress_callback: progress_callback(processed_paragraph_count, total_paragraphs_to_process)
            continue
            
        logger.info(f"Translating batch {batch_idx+1}/{total_api_batches} ({len(original_texts_in_batch)} paragraphs)")
        doc_translation_metrics.log_api_call() # Log one call for the batch to any engine

        translated_texts_in_batch = []
        try:
            if engine.lower() == 'gemini':
                translated_texts_in_batch = translation_using_gemini(
                    original_texts_in_batch, target_language, tone
                )
            elif engine.lower() == 'groq':
                translated_texts_in_batch = translation_using_groq(
                    original_texts_in_batch, target_language, tone
                )
            else:
                logger.error(f"Unsupported translation engine: {engine}. Skipping batch.")
                translated_texts_in_batch = original_texts_in_batch # Fallback to original
        except RuntimeError as e:
            logger.error(f"Translation aborted: {e}")
            if progress_callback:
                progress_callback(total_paragraphs_to_process, total_paragraphs_to_process)
            logger.error("Aborting translation and not saving output file due to missing API key.")
            return False  # Abort translation, do not save output file

        for para_obj, original_text, translated_text in zip(para_batch_list, original_texts_in_batch, translated_texts_in_batch):
            if not translated_text or (not translated_text.strip() and original_text.strip()):
                logger.warning(f"Empty or blank translation for: '{original_text[:50]}...'. Using original.")
                final_translated_text = original_text # Use original if translation is effectively empty
            else:
                final_translated_text = process_translated_text(original_text, translated_text)

            para_is_rtl = force_rtl or target_is_rtl_lang or get_paragraph_direction(final_translated_text)
            pPr = para_obj._element.get_or_add_pPr()
            bidi_elements = pPr.xpath('./w:bidi')
            if para_is_rtl:
                if not bidi_elements: bidi_tag = OxmlElement('w:bidi'); pPr.append(bidi_tag)
                else: bidi_tag = bidi_elements[0]
                bidi_tag.set(qn('w:val'), "true")
            else: 
                if bidi_elements: pPr.remove(bidi_elements[0])

            original_runs_data = all_paragraphs_with_runs.get(para_obj, [])
            if not original_runs_data: 
                para_obj.text = final_translated_text
            elif len(original_runs_data) == 1:
                first_run_obj = original_runs_data[0][0]; first_run_obj.text = final_translated_text
                for r_idx in range(1, len(para_obj.runs)): para_obj.runs[r_idx].text = ""
            else: 
                for r, _ in original_runs_data: r.text = "" # Clear runs
                is_list_item = is_numbered_list_paragraph(para_obj) or extract_list_marker(original_text) is not None
                if is_list_item or len(final_translated_text) < 30:
                    original_runs_data[0][0].text = final_translated_text
                else:
                    total_orig_len = sum(len(orig_rtxt) for _, orig_rtxt in original_runs_data)
                    if total_orig_len == 0: original_runs_data[0][0].text = final_translated_text
                    else:
                        current_pos = 0
                        for r_idx, (run_obj, orig_run_text) in enumerate(original_runs_data):
                            if current_pos >= len(final_translated_text): break
                            if r_idx == len(original_runs_data) - 1: run_obj.text = final_translated_text[current_pos:]
                            else:
                                proportion = len(orig_run_text) / total_orig_len
                                length_for_this_run = int(proportion * len(final_translated_text))
                                segment = final_translated_text[current_pos : current_pos + length_for_this_run]
                                if current_pos + length_for_this_run < len(final_translated_text):
                                    space_pos = segment.rfind(' '); 
                                    if space_pos > 0 : segment = segment[:space_pos]
                                run_obj.text = segment
                                current_pos += len(segment)
                                if current_pos < len(final_translated_text) and final_translated_text[current_pos] == ' ':
                                    current_pos += 1 
            
            for run_obj in para_obj.runs:
                if run_obj.text.strip():
                    run_is_rtl = force_rtl or target_is_rtl_lang or get_paragraph_direction(run_obj.text)
                    rPr = run_obj.element.get_or_add_rPr()
                    rtl_elements = rPr.xpath('./w:rtl')
                    if run_is_rtl:
                        if not rtl_elements: rtl_tag = OxmlElement('w:rtl'); rPr.append(rtl_tag)
                        else: rtl_tag = rtl_elements[0]
                        rtl_tag.set(qn('w:val'), "true")
                    else: 
                        if rtl_elements: rPr.remove(rtl_elements[0])
            
            processed_paragraph_count += 1
            if progress_callback:
                progress_callback(processed_paragraph_count, total_paragraphs_to_process)

    logger.info(f"Saving translated document to: {output_path}")
    try:
        doc.save(output_path)
        logger.info("Document saved successfully.")
        if progress_callback: progress_callback(total_paragraphs_to_process, total_paragraphs_to_process)
    except Exception as e:
        logger.error(f"Failed to save the processed document: {e}", exc_info=True)
        return False
    logger.info(f"Translation complete. Total API calls this session: {doc_translation_metrics.get_total_api_calls_session()}")
    return True

if __name__ == "__main__":
    import argparse
    import sys # Ensure sys is imported for sys.stdout, sys.exit

    parser = argparse.ArgumentParser(description="Translate DOCX files using specified engine.")
    parser.add_argument("input_file", help="Input DOCX file path")
    parser.add_argument("--output_file", help="Output DOCX file path (default: input_translated_lang.docx)")
    parser.add_argument("--target_language", default="hi", help="Target language code (e.g., 'es', 'fr', 'hi')")
    parser.add_argument("--engine", default="gemini", choices=['gemini', 'groq'], help="Translation engine to use.")
    parser.add_argument("--tone", default="professional", help="Tone for translation (e.g., 'professional', 'casual')")
    parser.add_argument("--force_rtl", action="store_true", help="Force right-to-left text direction regardless of language")
    parser.add_argument("--gemini_api_key", help="Google Gemini API Key (reads from GEMINI_API_KEY env var if not set)")
    parser.add_argument("--groq_api_key", help="Groq API Key (reads from GROQ_API_KEY env var if not set)")

    args = parser.parse_args()

    if args.gemini_api_key: os.environ["GEMINI_API_KEY"] = args.gemini_api_key
    if args.groq_api_key: os.environ["GROQ_API_KEY"] = args.groq_api_key
    
    if args.engine == "gemini" and not os.environ.get("GEMINI_API_KEY"):
        print("Error: Gemini engine selected but GEMINI_API_KEY not provided via argument or environment variable.")
        sys.exit(1)
    if args.engine == "groq" and not os.environ.get("GROQ_API_KEY"):
        print("Error: Groq engine selected but GROQ_API_KEY not provided via argument or environment variable.")
        sys.exit(1)

    if not args.output_file:
        base, ext = os.path.splitext(args.input_file)
        args.output_file = f"{base}_translated_{args.target_language}{ext}"

    # Setup detailed logging for CLI
    # logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(name)s: %(message)s")
    # Ensure logger for "advanced_docx_translator" is also at INFO or DEBUG if desired
    logging.getLogger("advanced_docx_translator").setLevel(logging.INFO)


    start_time = time.time()
    processed_paragraphs_count_for_progress = 0 # To allow progress to print final newline

    try:
        def cli_progress(current, total):
            global processed_paragraphs_count_for_progress
            processed_paragraphs_count_for_progress = total if current >=total else 0

            if total > 0:
                 percent = min((current / total) * 100, 100)
                 sys.stdout.write(f"\rProgress: {current}/{total} ({percent:.1f}%)")
                 sys.stdout.flush()
            else:
                 sys.stdout.write(f"\rProcessing item: {current}...")
                 sys.stdout.flush()
            if current >= total and total > 0:
                 sys.stdout.write("\n"); sys.stdout.flush()

        success = translate_docx_advanced(
            input_path=args.input_file,
            output_path=args.output_file,
            target_language=args.target_language,
            engine=args.engine,
            tone=args.tone,
            force_rtl=args.force_rtl,
            progress_callback=cli_progress
        )
        
        if processed_paragraphs_count_for_progress == 0 and "--output_file" in sys.argv : # If progress didn't complete but ran
             print()


        end_time = time.time()
        logger.info(f"Translation finished in {end_time - start_time:.2f} seconds.")

    except Exception as e:
        print() 
        logger.error(f"A critical error occurred during the translation process: {e}", exc_info=True)
        sys.exit(1)